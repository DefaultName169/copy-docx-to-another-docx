import argparse
import re
import shutil
import zipfile
import random
import zipfile
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.run import WD_BREAK
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

parser = argparse.ArgumentParser()
parser.add_argument('--path', default=None, required=False)
# parser.add_argument('--name', default=None, required=False)
cli_args = parser.parse_args()

def get_input(var_name):
    if auto_input := getattr(cli_args, var_name, None):
        print("Auto input:", auto_input)
        return auto_input
    else:
        return input("Manual input: ")

path = get_input("path")
# name  = get_input("name")

path = open(path, 'r', encoding='utf-8')
lines = path.readlines()
path.close()
pathname = re.sub('^__path__ = |\n', '' , lines[0])
output_name = lines[1]
output_name = re.sub(r'^name_of_output = |\n','',output_name)

shutil.copyfile('mau.docx', output_name + '.docx')
doc = docx.Document(output_name + '.docx')

num_of_figures = 1
num_of_tables = 1

def xmlstring_to_oxmlelement(string):
    queue = []
    save = []
    eval = ''
    first = ()
    array = re.split('<|>', string)
    for x in array:
        if x == '':
            continue
        else :
            my_arr = re.split(' \w+:',x)
            if queue:
                if re.search('^/' + queue[-1], x) :
                    a = queue.pop()
                    while save:
                        eval += a + '.append(' + save.pop(0) + ')' + '\n'
                    if queue:
                        eval += queue[-1] + '.append(' + a + ')' + '\n'
                    continue

                if queue[-1] == 't':
                    eval += 't.text = \'' + x + '\'' + '\n'
                    continue

            pin = re.sub('/', '' , my_arr[0])
            queue.append(pin)
            first = queue[0]
            eval += pin + ' = OxmlElement(\'w:'+pin+'\')'+'\n'
            i = 1
            while(i < len(my_arr)):
                my_value = re.split('="|"|/', my_arr[i])
                if my_value[1] == 'preserve' :
                    eval += my_arr[0] + '.set(qn(\'xml:' + my_value[0] + '\'),\'' + my_value[1] + '\')' +'\n'
                else:
                    eval += my_arr[0] + '.set(qn(\'w:' + my_value[0] + '\'),\'' + my_value[1] + '\')' +'\n'
                i += 1
            if re.search('/', my_arr[-1]) :
                pin = re.sub('/', '' , my_arr[0])
                save.append(pin)
                queue.pop()
    eval = 'output = ()\n' + eval + 'output = '+ first
    loc = {}
    exec(eval, globals(), loc)
    output = loc['output']
    return output


def random_id ():
    return ''.join(random.choices('0123456789ABCDEF', k=8))

def copy_docx(path) :
    string = ()
    with zipfile.ZipFile(path, 'r') as zip:
        string = zip.read('word/document.xml').decode('utf8')
    string = str(string)
    xml_string = re.finditer('<w:tbl>.+?</w:tbl>|<w:p.+?</w:p>' , string)
    for xml in xml_string :
        xml = xml.group()
        xml = re.sub('<w:', '<', xml)
        xml = re.sub('</w:', '</', xml)
        output = xmlstring_to_oxmlelement(xml)    
        doc.element.body[-2].addnext(output)

def add_text_to_docx(path) :
    copy_docx(path)

def add_picture_to_docx(path , name_picture) :
    global num_of_figures
    id_table_of_figures = random_id()
    doc.add_picture(path, width=Inches(7.5))
    paragraph = doc.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph._p.set(qn('w:rsidRDefault'),id_table_of_figures)
    if name_picture != '':
        xml_string = '<p w:rsidR="' + id_table_of_figures +'" w:rsidRDefault="'+ id_table_of_figures +'" w:rsidP="'+ id_table_of_figures +'"><pPr><pStyle w:val="Caption"/></pPr><r><t xml:space="preserve">Figure </t></r><fldSimple w:instr=" SEQ Figure \* ARABIC "><r><rPr><noProof/></rPr><t> ' + str(num_of_figures) + ' </t></r></fldSimple><r><t>: '+ name_picture +'</t></r></p>'
        output = xmlstring_to_oxmlelement(xml_string)    
        doc.paragraphs[-1]._p.addnext(output)


def add_table_to_docx(path, name_table) :
    global num_of_tables
    id_table_of_tables = random_id()
    copy_docx(path)
    xml_string = '<p w:rsidR="'+ id_table_of_tables +'" w:rsidRDefault="'+ id_table_of_tables +'" w:rsidP="'+ id_table_of_tables +'"><pPr><pStyle w:val="Caption"/><keepNext/><jc w:val="center"/></pPr><proofErr w:type="gramStart"/><r><t xml:space="preserve">Table </t></r><fldSimple w:instr=" SEQ Table \* ARABIC "><r><rPr><noProof/></rPr><t> '+ str(num_of_tables) +' </t></r></fldSimple><r><t>.</t></r><proofErr w:type="gramEnd"/><r><t xml:space="preserve"> '+ name_table +'</t></r></p>'
    num_of_tables += 1
    doc.paragraphs[-1]._p.set(qn('w:rsidRDefault'),id_table_of_tables)
    output = xmlstring_to_oxmlelement(xml_string)    
    doc.paragraphs[-2]._p.addnext(output)

lastlevel = -1
folder = pathname
tree_line_start = 0

for n in range(2 , len(lines)):
    if lines[n] == 'TREE :\n' :
        tree_line_start = n + 1
        break

for i in range(tree_line_start , len(lines)):
    if lines[i] == '\n' :
        continue
    else :
        level = len(re.findall('\t|\s{4}', lines[i])) + 1
        if level <= lastlevel : 
            pathname = re.sub(r'(.*)(/.*){'+ str(lastlevel - level + 1) + '}', r'\1', pathname)
            # for j in range(0, lastlevel - level + 1):
            #     pathname = re.sub('(.*)/.*', r'\1' , pathname)

        if level < lastlevel or level == 1:
            run = doc.paragraphs[-1].add_run()
            run.add_break(WD_BREAK.PAGE)
            # doc.add_page_break()
        
        name = re.sub('\n|\t|\s{4}', '' , lines[i])
        names = name.split(' : ')
        folder_name = names[0]
        name_header = ''
        if len(names) > 1 :
            name_header = names[1]
        pathname = pathname + '/' + folder_name
        print(pathname)

        isFolder = not re.search('.docx|.png|.jpg', folder_name)
        if isFolder :            
            if re.search(r'^\d+(\.(\d+))*', name_header) :
                name_header = re.sub(r'^\d+(\.(\d+))*\.(\s)?', '', name_header)
                new_p = doc.add_heading(name_header, level)
            else :
                new_p = doc.add_heading(name_header, level)
                numPr = OxmlElement('w:numPR')
                ilvl = OxmlElement('w:ilvl')
                ilvl.set(qn('w:val'),'0')
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'),'0')
                numPr.append(ilvl)
                numPr.append(numId)
                new_p._element.pPr.append(numPr)
                new_p._element.pPr.jc_val = WD_ALIGN_PARAGRAPH.CENTER

        elif re.search('.docx', folder_name) and name_header != '':

            add_table_to_docx(pathname , name_header)

        elif re.search('.docx',folder_name):

            add_text_to_docx(pathname)

        elif re.search('.jpg|.png', folder_name) :
            
            add_picture_to_docx(pathname, name_header)
        
        lastlevel = level

# def ReadAllDir(path, level) : 
#     dir_path = os.listdir(path)
#     i = 0
#     for folder_name in dir_path :
#         print(path + '/' + folder_name)
#         isFolder = not re.search('.docx|.png|.jpg', folder_name)
#         if isFolder :            
#             if re.search("\d+(.\d+)?(.\d+)?(\s+|.)?", folder_name) :
#                 string = re.sub("\d+(.\d+)?(.\d+)?(\s+|.)?", '', folder_name)
#                 new_p = doc.add_heading(string,level)
#             else :
#                 string = re.sub(r' ([A-Z]) ', r' \1 : ', folder_name)
#                 new_p = doc.add_heading(string,level)
#                 numPr = OxmlElement('w:numPR')
#                 ilvl = OxmlElement('w:ilvl')
#                 ilvl.set(qn('w:val'),'0')
#                 numId = OxmlElement('w:numId')
#                 numId.set(qn('w:val'),'0')
#                 numPr.append(ilvl)
#                 numPr.append(numId)
#                 new_p._element.pPr.append(numPr)
#                 new_p._element.pPr.jc_val = WD_ALIGN_PARAGRAPH.CENTER

#             ReadAllDir(path + '/' + folder_name,level + 1)

#         elif re.search('Table.*\.docx', folder_name):

#             add_table_to_docx(path + '/' + folder_name)

#         elif re.search('.docx',folder_name):

#             add_text_to_docx(path + '/' + folder_name)

#         elif re.search('.jpg|.png', folder_name) :

#             add_picture_to_docx(path + '/' + folder_name)

#         if( (level == 1 or level == 3) and i < len(dir_path)) :
#             try :
#                 if not doc.paragraphs[-1].runs[-1]._element.br_lst :
#                     doc.add_page_break()
#             except:
#                 doc.add_page_break()
#         i += 1

# ReadAllDir(path,1)

doc.save(output_name + '.docx')